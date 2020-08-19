import docx  #hay que usar es python-docx y se invoca con docx
import datetime
import time
from open_close_program import ifOpenThenClose, isOpen, findWindowProgram, clickAndSave


def getConfig(f_path):
     config = []
     with open(f_path, "rt") as t_file:
          for each in t_file:
               config.append(each.replace("\n",""))
     return config


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
     title = 'F603.docx' + ' - Word'
     config = getConfig('config.txt')
     filePath = config[0]
     docName = config[1]

     if isOpen(title) == True:
          x, y = findWindowProgram(title)
          clickAndSave(title)
          ifOpenThenClose(filePath, title)
          time.sleep(1)


     doc = docx.Document(filePath)
     paragraphs  = doc.paragraphs
     text = paragraphs[0].text
     paragraphs[0]._p.clear()

     paragraphs[0].add_run('------------1---------> '+str(datetime.datetime.now().today().strftime('%Y-%m-%d  time: %H:%M:%S'))+' \n')
     paragraphs[0].add_run('\n')
     paragraphs[0].add_run(' Nueva entrada de texto!! \n')
     paragraphs[0].add_run('\n')
     paragraphs[0].add_run('\n')
     paragraphs[0].add_run('\n')
     paragraphs[0].add_run('\n')
     #paragraphs[1].add_run('-----------------------------------< \n')
     paragraphs[0].add_run(text)
     doc.save(filePath)

     if isOpen(title) == False:
          ifOpenThenClose(filePath,title)
